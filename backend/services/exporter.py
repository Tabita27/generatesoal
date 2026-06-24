from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io

def export_to_docx(judul: str, questions: list) -> bytes:
    doc = Document()

    # Judul
    title = doc.add_heading(judul, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2D, 0x3A, 0x8C)

    doc.add_paragraph(f"Total soal: {len(questions)}")
    doc.add_paragraph("─" * 60)

    for q in questions:
        q_dict = q if isinstance(q, dict) else q.dict()
        nomor = q_dict.get("nomor", "")
        pertanyaan = q_dict.get("pertanyaan", "")
        tipe = q_dict.get("tipe", "")
        pilihan = q_dict.get("pilihan", {})
        kunci = q_dict.get("kunci_jawaban", "")
        kesulitan = q_dict.get("tingkat_kesulitan", "")

        # Pertanyaan
        p = doc.add_paragraph()
        run = p.add_run(f"{nomor}. [{tipe.replace('_', ' ').title()}] [{kesulitan.upper()}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)

        doc.add_paragraph(pertanyaan)

        # Pilihan (jika PG)
        if pilihan and isinstance(pilihan, dict):
            for huruf, isi in pilihan.items():
                doc.add_paragraph(f"   {huruf}. {isi}", style="List Bullet")

        # Kunci jawaban
        kunci_p = doc.add_paragraph()
        kunci_run = kunci_p.add_run(f"✓ Kunci: {kunci}")
        kunci_run.font.color.rgb = RGBColor(0x16, 0x7A, 0x3B)
        kunci_run.bold = True

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def export_to_pdf(judul: str, questions: list) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    story = []

    # Judul
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        textColor='#2D3A8C',
        spaceAfter=12
    )
    story.append(Paragraph(judul, title_style))
    story.append(Paragraph(f"Total soal: {len(questions)}", styles['Normal']))
    story.append(Spacer(1, 0.5*cm))

    question_style = ParagraphStyle(
        'QuestionStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leading=16
    )

    for q in questions:
        q_dict = q if isinstance(q, dict) else q.dict()
        nomor = q_dict.get("nomor", "")
        pertanyaan = q_dict.get("pertanyaan", "")
        tipe = q_dict.get("tipe", "")
        pilihan = q_dict.get("pilihan", {})
        kunci = q_dict.get("kunci_jawaban", "")
        kesulitan = q_dict.get("tingkat_kesulitan", "")

        header = f"<b>{nomor}. [{tipe.replace('_', ' ').title()}] [{kesulitan.upper()}]</b>"
        story.append(Paragraph(header, question_style))
        story.append(Paragraph(pertanyaan, question_style))

        if pilihan and isinstance(pilihan, dict):
            for huruf, isi in pilihan.items():
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{huruf}. {isi}", question_style))

        kunci_style = ParagraphStyle(
            'KunciStyle',
            parent=styles['Normal'],
            fontSize=11,
            textColor='#167A3B',
            spaceAfter=10
        )
        story.append(Paragraph(f"<b>✓ Kunci: {kunci}</b>", kunci_style))
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    return buffer.getvalue()